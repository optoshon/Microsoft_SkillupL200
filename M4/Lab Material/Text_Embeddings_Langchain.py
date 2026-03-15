# Text_Embeddings_Langchain.py
# ============================================================================
# MODULE 3: TEXT INTELLIGENCE ENGINE
# ============================================================================
#
# PURPOSE:
#   Load extracted DOCX files (from Module 1)
#   Apply recursive character-based chunking to preserve semantics
#   Generate stable IDs for tracking and updates
#   Create embeddings using Azure OpenAI
#   Store in Chroma vector database
#
# WHAT IT DOES:
#   1) Reads DOCX files from docling_output/
#   2) Applies recursive character-based text splitting
#   3) Generates stable chunk IDs for auditing/tracking
#   4) Embeds all chunks using Azure OpenAI embeddings API
#   5) Stores in Chroma vector database with metadata
#
# INPUT REQUIRED:
#   - docling_output/*/text/*.docx (from Module 1: Read_File_Docling.py)
#   - .env file with Azure OpenAI embeddings deployment
#
# OUTPUT CREATED:
#   - vector_db_text_chroma/ (Chroma database with text embeddings)
#
# RUNTIME: 30-120 seconds (depends on document size)
#
# PREREQUISITE:
#   Run Module 1 (Read_File_Docling.py) BEFORE this script
#
# ============================================================================
# SETUP INSTRUCTIONS
# ============================================================================
#
# Step 1: ACTIVATE PYTHON ENVIRONMENT
#   conda activate multimodal_rag
#   # OR
#   source myenv310/bin/activate  # (macOS/Linux)
#   myenv310\Scripts\activate      # (Windows)
#
# Step 2: INSTALL DEPENDENCIES
#   If you get import errors, uncomment and run:

# !pip install -q python-dotenv python-docx langchain-core langchain-text-splitters langchain-openai chromadb langchain-chroma
# OR for full requirements:
# !pip install -r requirements.txt

# Step 3: ENSURE MODULE 1 COMPLETED
#   Run: python Read_File_Docling.py
#   Verify: Check that docling_output/ folder exists with extracted text
#
# Step 4: CONFIGURE PATHS & AZURE CREDENTIALS
#   - Search for "PUT YOUR PATH HERE" below
#   - Create/update .env file with Azure OpenAI credentials
#   - Verify AZURE_OPENAI_EMBEDDINGS_DEPLOYMENT deployment name
#
# Step 5: RUN THE SCRIPT
#   python Text_Embeddings_Langchain.py
#
# ============================================================================

# =========================
# 0) Install (run once)
# =========================
# !pip install -q python-dotenv python-docx langchain-core langchain-text-splitters langchain-openai chromadb langchain-chroma

# =========================
# 1) Config (edit only here)
# =========================
from pathlib import Path

# PUT YOUR PATH HERE: Where Module 1 (Read_File_Docling.py) saved extracted documents
DOCLING_OUTPUT_DIR = Path("C:\\Users\\srika\\Documents\\Microsoft_19jan\\MS_NOIDA_FEB\\M4 packet\\docling_output")

# PUT YOUR PATH HERE: Where to save Chroma vector database for text (separate from table/image DBs)
PERSIST_DIR = Path("C:\\Users\\srika\\Documents\\Microsoft_19jan\\MS_NOIDA_FEB\\M4 packet\\vector_db_text_chroma")

COLLECTION_NAME = "text_chunks_v1"

# Chunking parameters for text
CHUNK_SIZE = 1200
CHUNK_OVERLAP = 150
SEPARATORS = ["\n\n", "\n", ". ", " ", ""]

# Optional: to limit to exactly 10 docs for testing, set to 10; else None
MAX_DOCS = None

# =========================
# 2) Load .env (NO FALLBACKS)
# =========================
import os
from dotenv import load_dotenv

# PUT YOUR PATH HERE: Location of .env file with Azure OpenAI credentials
load_dotenv(dotenv_path=Path(".env"), override=False)

REQUIRED_ENV_VARS = [
    "AZURE_OPENAI_ENDPOINT",
    "AZURE_OPENAI_API_KEY",
    "AZURE_OPENAI_API_VERSION",
    "AZURE_OPENAI_EMBEDDINGS_DEPLOYMENT",
]

missing = [k for k in REQUIRED_ENV_VARS if not os.getenv(k)]
if missing:
    raise RuntimeError(
        "Missing required .env variables (no fallbacks allowed):\n"
        + "\n".join(f"  - {k}" for k in missing)
        + "\n\nExpected .env example:\n"
          "AZURE_OPENAI_ENDPOINT=https://<your-resource>.openai.azure.com/\n"
          "AZURE_OPENAI_API_KEY=<your-key>\n"
          "AZURE_OPENAI_API_VERSION=2024-02-15-preview\n"
          "AZURE_OPENAI_EMBEDDINGS_DEPLOYMENT=<your-embeddings-deployment-name>\n"
    )

AZURE_OPENAI_ENDPOINT = os.environ["AZURE_OPENAI_ENDPOINT"]
AZURE_OPENAI_API_KEY = os.environ["AZURE_OPENAI_API_KEY"]
AZURE_OPENAI_API_VERSION = os.environ["AZURE_OPENAI_API_VERSION"]
AZURE_OPENAI_EMBEDDINGS_DEPLOYMENT = os.environ["AZURE_OPENAI_EMBEDDINGS_DEPLOYMENT"]

# =========================
# 3) Read Docling-extracted DOCX files
# =========================
from docx import Document as DocxDocument
from langchain_core.documents import Document as LCDocument

def read_docx_text(docx_path: Path) -> str:
    doc = DocxDocument(docx_path)
    parts = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)
    return "\n\n".join(parts).strip()

def find_docling_docx_files(root: Path) -> list[Path]:
    # Expected structure: docling_output/<doc_id>/text/<doc_id>.docx
    return sorted(root.glob("*/text/*.docx"))

docx_files = find_docling_docx_files(DOCLING_OUTPUT_DIR)
if not docx_files:
    raise RuntimeError(
        f"No .docx files found under {DOCLING_OUTPUT_DIR}.\n"
        "Expected: docling_output/<doc_id>/text/<doc_id>.docx"
    )

if MAX_DOCS is not None:
    docx_files = docx_files[:MAX_DOCS]

base_docs: list[LCDocument] = []
for p in docx_files:
    doc_id = p.parent.parent.name  # folder above /text
    text = read_docx_text(p)
    if not text:
        continue
    base_docs.append(
        LCDocument(
            page_content=text,
            metadata={
                "doc_id": doc_id,
                "source_docx": str(p),
                "modality": "text",
                "extracted_by": "docling",
            },
        )
    )

if not base_docs:
    raise RuntimeError("All DOCX files were empty after reading. Nothing to index.")

print(f"Loaded {len(base_docs)} documents from Docling output.")

# =========================
# 4) Chunk using LangChain
# =========================
from langchain_text_splitters import RecursiveCharacterTextSplitter

splitter = RecursiveCharacterTextSplitter(
    chunk_size=CHUNK_SIZE,
    chunk_overlap=CHUNK_OVERLAP,
    separators=SEPARATORS,
)

chunked_docs = splitter.split_documents(base_docs)
if not chunked_docs:
    raise RuntimeError("Chunking produced 0 chunks. Check your inputs/config.")

# Add stable IDs (useful later for upserts / tracking)
import hashlib

def stable_chunk_id(d: LCDocument, chunk_idx: int) -> str:
    base = f"{d.metadata.get('doc_id','')}::{d.metadata.get('source_docx','')}::{chunk_idx}::{d.page_content}"
    return hashlib.sha256(base.encode("utf-8")).hexdigest()[:24]

for i, d in enumerate(chunked_docs, start=1):
    d.metadata["chunk_index_global"] = i
    # per-doc chunk index can be derived later; keep a stable id now
    d.metadata["chunk_uid"] = stable_chunk_id(d, i)

print(f"Created {len(chunked_docs)} chunks.")

# =========================
# 5) Embeddings (Azure OpenAI) from .env
# =========================
from langchain_openai import AzureOpenAIEmbeddings

embeddings = AzureOpenAIEmbeddings(
    azure_endpoint=AZURE_OPENAI_ENDPOINT,
    api_key=AZURE_OPENAI_API_KEY,
    api_version=AZURE_OPENAI_API_VERSION,
    azure_deployment=AZURE_OPENAI_EMBEDDINGS_DEPLOYMENT,
)

# =========================
# 6) Persist to a local Chroma vector DB
# =========================
from langchain_chroma import Chroma

PERSIST_DIR.mkdir(parents=True, exist_ok=True)

# Chroma uses document IDs optionally; we'll pass our stable IDs
ids = [d.metadata["chunk_uid"] for d in chunked_docs]

vectorstore = Chroma(
    collection_name=COLLECTION_NAME,
    embedding_function=embeddings,
    persist_directory=str(PERSIST_DIR),
)

# Upsert (add_documents will insert; Chroma will generally ignore duplicates by id if same)
vectorstore.add_documents(documents=chunked_docs, ids=ids)

print(f"✅ Stored {len(chunked_docs)} chunk vectors in Chroma.")
print(f"   Persist dir: {PERSIST_DIR.resolve()}")
print(f"   Collection:  {COLLECTION_NAME}")

# =========================
# 7) Quick sanity retrieval test (no LLM, just similarity)
# =========================
query = "test query: policy objective"
hits = vectorstore.similarity_search(query, k=3)

print("\nTop-3 hits:")
for j, h in enumerate(hits, start=1):
    print(f"\n--- Hit {j} ---")
    print("metadata:", {k: h.metadata[k] for k in ["doc_id", "chunk_uid", "source_docx"] if k in h.metadata})
    print("text preview:", h.page_content[:300].replace("\n", " "))
