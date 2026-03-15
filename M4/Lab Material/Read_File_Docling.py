# Read_File_Docling.py
# ============================================================================
# MODULE 1: DOCUMENT EXTRACTION
# ============================================================================
#
# PURPOSE:
#   Extract text, tables, and images from enterprise documents (PDF, DOCX, PPTX, HTML)
#   Outputs organized into docling_output/ folder for processing by other modules
#
# WHAT IT DOES:
#   1) Text   → Saved as DOCX files per document
#   2) Images → Saved as PNG files per figure
#   3) Tables → Saved as CSV + SQLite database
#   4) Creates manifest.json metadata index
#
# INPUT REQUIRED:
#   - Documents in INPUT_DIR folder (PDF, DOCX, PPTX, HTML, MD)
#   - No Azure credentials needed for this module
#
# OUTPUT CREATED:
#   - docling_output/<doc_id>/text/<doc_id>.docx
#   - docling_output/<doc_id>/tables/table_001.csv
#   - docling_output/<doc_id>/images/figure_001.png
#   - docling_output/tables.sqlite
#   - docling_output/manifest.json
#
# RUNTIME: 30-120 seconds per document (depends on complexity)
#
# ============================================================================
# SETUP INSTRUCTIONS
# ============================================================================
#
# Step 1: CREATE PYTHON ENVIRONMENT
#   Option A (Conda):
#     conda create -n multimodal_rag python=3.10
#     conda activate multimodal_rag
#
#   Option B (venv):
#     python -m venv multimodal_rag
#     multimodal_rag\Scripts\activate  # Windows
#     source multimodal_rag/bin/activate  # macOS/Linux
#
#   Option C (Existing env):
#     conda activate myenv310  # or your env name
#
# Step 2: INSTALL DEPENDENCIES
#   If you get import errors, uncomment and run the pip install below:

# !pip install docling docling-core pandas pillow python-docx
# OR for full requirements:
# !pip install -r requirements.txt

# Step 3: CONFIGURE PATHS
#   Search for "PUT YOUR PATH HERE" below and update INPUT_DIR and OUTPUT_DIR
#
# Step 4: RUN THE SCRIPT
#   python Read_File_Docling.py
#
# ============================================================================

from pathlib import Path
import re
import json
import logging
import sqlite3

import pandas as pd
from docx import Document as DocxDocument

from docling.document_converter import DocumentConverter, PdfFormatOption
from docling.datamodel.base_models import InputFormat
from docling.datamodel.pipeline_options import PdfPipelineOptions
from docling_core.types.doc import TextItem, PictureItem

# =========================
# CONFIG (edit only here)
# =========================

# PUT YOUR PATH HERE: Replace with folder containing your documents (PDF, DOCX, PPTX, HTML)
INPUT_DIR = Path(r"C:\Users\shonr\OneDrive - Tekframeworks\Training\Microsoft\Microsoft_SkillupL200\M4\Lab Material")

# PUT YOUR PATH HERE: Where extracted artifacts will be saved (can be relative or absolute)
OUTPUT_DIR = Path("docling_output")

SQLITE_NAME = "tables.sqlite"
PDF_IMAGE_SCALE = 2.0               # Higher = sharper figure images

# File types to process
SUPPORTED_EXTS = {".pdf", ".docx", ".pptx", ".html", ".htm", ".md"}

# =========================


logging.basicConfig(level=logging.INFO)
LOG = logging.getLogger("docling")


def safe_name(s: str) -> str:
    return re.sub(r"[^A-Za-z0-9._-]+", "_", s).strip("_") or "untitled"


def build_converter():
    pdf_opts = PdfPipelineOptions()
    pdf_opts.images_scale = PDF_IMAGE_SCALE
    pdf_opts.generate_page_images = False
    pdf_opts.generate_picture_images = True

    return DocumentConverter(
        format_options={
            InputFormat.PDF: PdfFormatOption(pipeline_options=pdf_opts)
        }
    )


def collect_text_blocks(conv_res):
    blocks = []
    for element, _ in conv_res.document.iterate_items():
        if isinstance(element, TextItem):
            txt = (element.text or "").strip()
            if txt:
                blocks.append(txt)
    return blocks


def write_docx(text_blocks, out_path: Path):
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc = DocxDocument()
    for block in text_blocks:
        doc.add_paragraph(block)
    doc.save(out_path)


def export_tables(conv_res, tables_dir: Path, sqlite_conn, doc_id: str):
    tables_dir.mkdir(parents=True, exist_ok=True)
    meta = []

    for i, tbl in enumerate(conv_res.document.tables, start=1):
        df = tbl.export_to_dataframe(doc=conv_res.document)

        csv_path = tables_dir / f"table_{i:03d}.csv"
        df.to_csv(csv_path, index=False)

        sql_table = safe_name(f"{doc_id}__table_{i:03d}")
        df.to_sql(sql_table, sqlite_conn, if_exists="replace", index=False)

        meta.append({
            "table_index": i,
            "csv": str(csv_path),
            "sqlite_table": sql_table,
            "rows": int(df.shape[0]),
            "cols": int(df.shape[1]),
        })

    return meta


def export_images(conv_res, images_dir: Path):
    images_dir.mkdir(parents=True, exist_ok=True)
    meta = []
    counter = 0

    for element, _ in conv_res.document.iterate_items():
        if isinstance(element, PictureItem):
            counter += 1
            img_path = images_dir / f"figure_{counter:03d}.png"
            element.get_image(conv_res.document).save(img_path, "PNG")

            meta.append({
                "figure_index": counter,
                "path": str(img_path),
            })

    return meta


def main():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    files = [p for p in INPUT_DIR.iterdir()
             if p.is_file() and p.suffix.lower() in SUPPORTED_EXTS]

    if not files:
        raise RuntimeError(f"No supported documents found in {INPUT_DIR}")

    LOG.info("Found %d documents", len(files))

    converter = build_converter()
    sqlite_path = OUTPUT_DIR / SQLITE_NAME
    sqlite_conn = sqlite3.connect(sqlite_path)

    manifest = []

    try:
        for file_path in files:
            doc_id = safe_name(file_path.stem)
            doc_dir = OUTPUT_DIR / doc_id

            text_path = doc_dir / "text" / f"{doc_id}.docx"
            tables_dir = doc_dir / "tables"
            images_dir = doc_dir / "images"

            LOG.info("Processing: %s", file_path.name)
            conv_res = converter.convert(file_path)

            # Text
            text_blocks = collect_text_blocks(conv_res)
            write_docx(text_blocks, text_path)

            # Tables
            tables_meta = export_tables(conv_res, tables_dir, sqlite_conn, doc_id)

            # Images
            images_meta = export_images(conv_res, images_dir)

            record = {
                "source": str(file_path),
                "doc_id": doc_id,
                "text_docx": str(text_path),
                "tables_dir": str(tables_dir),
                "images_dir": str(images_dir),
                "sqlite_db": str(sqlite_path),
                "counts": {
                    "text_blocks": len(text_blocks),
                    "tables": len(tables_meta),
                    "images": len(images_meta),
                },
                "tables": tables_meta,
                "images": images_meta,
            }

            manifest.append(record)

            LOG.info(
                "Done %-25s | blocks=%d tables=%d images=%d",
                file_path.name,
                record["counts"]["text_blocks"],
                record["counts"]["tables"],
                record["counts"]["images"],
            )

    finally:
        sqlite_conn.commit()
        sqlite_conn.close()

    manifest_path = OUTPUT_DIR / "manifest.json"
    manifest_path.write_text(json.dumps(manifest, indent=2), encoding="utf-8")
    LOG.info("Wrote manifest: %s", manifest_path)
    LOG.info("SQLite DB: %s", sqlite_path)


if __name__ == "__main__":
    main()
